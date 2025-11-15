import calendar

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from typing import List, Optional, Dict, Any
from datetime import datetime
from configuraciones.configuracion import app_configuracion
from reportes.reporte_base import ReporteBase
from excepciones.base_datos_error import BaseDatosError

from repositorios.alumnos.alumno_repositorio import AlumnoRepositorio
from repositorios.alumnos.inscripcion_repositorio import InscripcionRepositorio
from repositorios.alumnos.info_clinica_alumno_repositorio import InfoClinicaAlumnoRepositorio
from repositorios.empleados.detalle_cargo_repositorio import DetalleCargoRepositorio
from repositorios.especialidades.especialidad_repositorio import EspecialidadRepositorio

from servicios.alumnos.alumno_servicio import AlumnoServicio
from servicios.alumnos.inscripcion_servicio import InscripcionServicio
from servicios.alumnos.info_clinica_alumno_servicio import InfoClinicaAlumnoServicio
from servicios.empleados.detalle_cargo_servicio import DetalleCargoServicio
from servicios.especialidades.especialidad_servicio import EspecialidadServicio


class ReporteInformeEducativoAlumnos(ReporteBase):
    def __init__(self):
        self.RUTA_REPORTES_INFORMES_EDUCATIVOS_ALUMNOS = app_configuracion.DIRECTORIO_REPORTES_ALUMNOS
    
    def transformar_lista_data_alumnos(
        self, alumno_repositorio: AlumnoRepositorio,
        inscripcion_repositorio: InscripcionRepositorio,
        info_clinica_alumno_repositorio: InfoClinicaAlumnoRepositorio,
        especialidad_id: int,
        lista_evaluadores: List
    ) -> List[Dict]:
        lista_dict_data_alumnos = []
        
        alumno_servicio = AlumnoServicio(alumno_repositorio)
        inscripcion_servicio = InscripcionServicio(inscripcion_repositorio)
        info_clinica_alumno_servicio = InfoClinicaAlumnoServicio(info_clinica_alumno_repositorio)
        
        for alumno in inscripcion_servicio.obtener_inscripcion_por_especialidad(especialidad_id):
            alumno_id = alumno[0]
            
            primer_nombre = alumno_servicio.obtener_alumno_por_id(alumno_id)[2]
            segundo_nombre = alumno_servicio.obtener_alumno_por_id(alumno_id)[3] if (alumno_servicio.obtener_alumno_por_id(alumno_id)[3] is not None) else ""
            tercer_nombre = alumno_servicio.obtener_alumno_por_id(alumno_id)[4] if (alumno_servicio.obtener_alumno_por_id(alumno_id)[4] is not None) else ""
            apellido_paterno = alumno_servicio.obtener_alumno_por_id(alumno_id)[5]
            apellido_materno = alumno_servicio.obtener_alumno_por_id(alumno_id)[6] if (alumno_servicio.obtener_alumno_por_id(alumno_id)[6] is not None) else ""
            nombre_completo = f"{primer_nombre} {segundo_nombre} {tercer_nombre} {apellido_paterno} {apellido_materno}"
            
            cedula_alumno = alumno[1]
            
            lugar_nacimiento = alumno_servicio.obtener_alumno_por_id(alumno_id)[9]
            fecha_nacimiento = alumno_servicio.obtener_alumno_por_id(alumno_id)[7]
            lugar_y_fecha_nacimiento = f"{lugar_nacimiento} {fecha_nacimiento}"
            
            edad = f"{alumno_servicio.obtener_alumno_por_id(alumno_id)[8]} años"
            sexo = "Masculino" if (alumno_servicio.obtener_alumno_por_id(alumno_id)[10] == "M") else "Femenino"
            
            procedencia = alumno_servicio.obtener_info_academica_alumno(alumno_id)[2]
            
            lista_diagnosticos = []
            lista_medicacion = []
            lista_certificados_discapacidad = []
            
            for diagnostico in info_clinica_alumno_servicio.obtener_info_clinica_por_alumno_id(alumno_id):
                nombre_diagnostico = diagnostico[2] if (diagnostico[2] is not None) else ""
                medicacion = diagnostico[7] if (diagnostico[7] is not None) else ""
                certificado_discapacidad = diagnostico[5] if (diagnostico[5] is not None) else ""
                medico_tratante = diagnostico[4] if (diagnostico[4] is not None) else ""
                fecha_diagnostico = diagnostico[3] if (diagnostico[3] is not None) else ""
                observacion_adicional = diagnostico[8] if (diagnostico[8] is not None) else ""
                
                lista_diagnosticos.append((nombre_diagnostico, medico_tratante, fecha_diagnostico, observacion_adicional))
                lista_medicacion.append((medicacion))
                lista_certificados_discapacidad.append((certificado_discapacidad))
            
            fecha_ingreso_taller = alumno_servicio.obtener_alumno_por_id(alumno_id)[13]
            fecha_ingreso_especialidad = alumno[6]
            
            nombre_representante = alumno_servicio.obtener_datos_representante(alumno_id)[3]
            apellido_representante = alumno_servicio.obtener_datos_representante(alumno_id)[4]
            nombre_completo_representante = f"{nombre_representante} {apellido_representante}"
            cedula_representante = alumno_servicio.obtener_datos_representante(alumno_id)[2]
            
            num_telefono_principal = alumno_servicio.obtener_datos_representante(alumno_id)[6]
            num_telefono_secundario = alumno_servicio.obtener_datos_representante(alumno_id)[7] if (alumno_servicio.obtener_datos_representante(alumno_id)[7] is not None) else ""
            lista_telefonos = [num_telefono_principal, num_telefono_secundario]
            
            if (len(lista_telefonos) == 2):
                telefonos = " ".join(lista_telefonos)
            else:
                telefonos = num_telefono_principal
            
            direccion_residencia = alumno_servicio.obtener_datos_representante(alumno_id)[5]
            
            especialidad_ocupacional = alumno[4]
            
            mes_evaluacion = int(datetime.now().date().month)
            anio_evaluacion = int(datetime.now().date().year)
            fecha_evaluacion = f"{mes_evaluacion}-{anio_evaluacion}"
            
            anio_escolar = f"{anio_evaluacion}-{anio_evaluacion + 1}"
            
            elemento_lista = {
                "nombre_completo_alumno": nombre_completo,
                "cedula_alumno": cedula_alumno,
                "lugar_y_fecha_nacimiento": lugar_y_fecha_nacimiento,
                "edad": edad,
                "sexo": sexo,
                "procedencia": procedencia,
                "lista_diagnosticos": lista_diagnosticos,
                "lista_medicacion": ", ".join(lista_medicacion),
                "lista_certificados_discapacidad": ", ".join(lista_certificados_discapacidad),
                "fecha_ingreso_taller": fecha_ingreso_taller,
                "fecha_ingreso_especialidad": fecha_ingreso_especialidad,
                "nombre_completo_representante": nombre_completo_representante,
                "cedula_representante": cedula_representante,
                "telefonos": telefonos,
                "direccion_residencia": direccion_residencia,
                "especialidad_ocupacional": especialidad_ocupacional,
                "lista_docentes_evaluadores": " - ".join(lista_evaluadores),
                "fecha_evaluacion": fecha_evaluacion,
                "anio_escolar": anio_escolar
            }
            
            lista_dict_data_alumnos.append(elemento_lista)
            
        return lista_dict_data_alumnos
    
    def transformar_lista_evaluadores(self, detalles_cargo_repositorio: DetalleCargoRepositorio, especialidad_id: int) -> List[str]:
        lista_evaluadores = []
        
        detalles_cargo_servicio = DetalleCargoServicio(detalles_cargo_repositorio)
        
        for evaluador in detalles_cargo_servicio.obtener_detalles_cargo_por_especialidad(especialidad_id):
            nombre_evaluador = evaluador[3]
            apellido_evaluador = evaluador[6]
            
            elemento_lista = f"Prof. {nombre_evaluador} {apellido_evaluador}"
            lista_evaluadores.append(elemento_lista)
        
        return lista_evaluadores
    
    def transformar_lista_firmantes(self, lista_evaluadores: List, nombre_completo_directora: str, nombre_completo_coord_academico: str) -> List[Dict]:
        lista_dict_firmantes = []
        
        lista_evaluadores.append(nombre_completo_directora)
        lista_evaluadores.append(nombre_completo_coord_academico)
        
        for firmante in lista_evaluadores:
            elemento_lista = {
                "nombre_completo_firmante": firmante
            }
            
            lista_dict_firmantes.append(elemento_lista)
        
        return lista_dict_firmantes
    
    def crear_documento(self):
        documento = Document()
        
        return documento
    
    def cargar_mes(self) -> str:
        fecha_actual = datetime.today()
        
        diccionario_meses_espaniol = {
            "January": "Enero",
            "February": "Febrero",
            "March": "Marzo",
            "April": "Abril",
            "May": "Mayo",
            "June": "Junio",
            "July": "Julio",
            "August": "Agosto",
            "September": "Septiembre",
            "October": "Octubre",
            "November": "Noviembre",
            "December": "Diciembre"
        }
        
        # Obtener el nombre del mes en inglés de la fecha de nacimiento
        nombre_mes_ingles = calendar.month_name[fecha_actual.month]
        
        # Devolver el nombre del mes en español
        return diccionario_meses_espaniol.get(nombre_mes_ingles)
    
    def cargar_cintillo(self, documento):
        encabezado = documento.sections[0].header

        encabezado_parrafo = encabezado.paragraphs[0] if encabezado.paragraphs else encabezado.add_paragraph()
        encabezado_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = encabezado_parrafo.add_run()
        run.add_picture("reportes/imagenes/CINTILLO_TELA.png")
    
    def cargar_fuente(self, run, nombre_fuente: str, tamanio: int):
        run.font.name = nombre_fuente
        run.font.size = Pt(tamanio)
    
    def cargar_espacio_firmas(self, documento, lista_dict_firmantes: List[Dict]):
        num_firmantes = len(lista_dict_firmantes)
        num_firmas_por_fila = 2
        num_filas = (num_firmantes + num_firmas_por_fila - 1) // num_firmas_por_fila
        
        parrafo = documento.add_paragraph()
        run = parrafo.add_run("Firmas Conformes:")
        run.bold = True
        self.cargar_fuente(run, "Arial", 12)
        
        self.agregar_salto_linea(documento, 1)
        
        tabla = documento.add_table(rows = num_filas, cols = num_firmas_por_fila, style = None)
        tabla.autofit = True
        
        indice_firmante = 0
        for fila in range(num_filas):
            fila_celdas = tabla.rows[fila].cells
            
            for columna in range(num_firmas_por_fila):
                if (indice_firmante < num_firmantes):
                    firmante = lista_dict_firmantes[indice_firmante]
                    nombre_firmante = firmante["nombre_completo_firmante"]
                    
                    celda = fila_celdas[columna]
                    
                    parrafo_linea = celda.add_paragraph("_________________________")
                    parrafo_linea.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    parrafo_nombre = celda.add_paragraph()
                    parrafo_nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = parrafo_nombre.add_run(f"{nombre_firmante}")
                    self.cargar_fuente(run, "Arial", 12)
                    
                    indice_firmante += 1
                else:
                    pass
    
    def agregar_salto_linea(self, documento, cantidad = 1):
        for _ in range(cantidad):
            parrafo = documento.add_paragraph(" ")
            parrafo.paragraph_format.space_after = Pt(12)
    
    def obtener_nombre_directora(self, detalles_cargo_repositorio: DetalleCargoRepositorio) -> str:
        detalles_cargo_servicio = DetalleCargoServicio(detalles_cargo_repositorio)
        
        for empleado in detalles_cargo_servicio.obtener_detalles_cargo_empleados():
            funcion_cargo = empleado[3]
            
            if (funcion_cargo == "DIRECTORA ENCARGADA"):
                empleado_id = empleado[0]
                
                nombre_directora = detalles_cargo_servicio.obtener_detalles_cargo_por_id(empleado_id)[2]
                apellido_directora = detalles_cargo_servicio.obtener_detalles_cargo_por_id(empleado_id)[5]
                nombre_completo_directora = f"{nombre_directora} {apellido_directora}"
                break
            else:
                continue
        
        return nombre_completo_directora
    
    def cargar_info_alumnos(self, documento, lista_dict_alumnos: List[Dict], lista_dict_firmantes: List[Dict]):
        for alumno in lista_dict_alumnos:
            nombre_completo_alumno = alumno["nombre_completo_alumno"]
            cedula_alumno = alumno["cedula_alumno"]
            lugar_y_fecha_nacimiento = alumno["lugar_y_fecha_nacimiento"]
            edad = alumno["edad"]
            sexo = alumno["sexo"]
            procedencia = alumno["procedencia"]
            diagnosticos_con_obervacion = alumno["lista_diagnosticos"]
            medicacion = alumno["lista_medicacion"]
            certificado_discapacidad = alumno["lista_certificados_discapacidad"]
            fecha_ingreso_taller = alumno["fecha_ingreso_taller"]
            fecha_ingreso_especialidad = alumno["fecha_ingreso_especialidad"]
            nombre_completo_representante = alumno["nombre_completo_representante"]
            cedula_representante = alumno["cedula_representante"]
            telefono_representante = alumno["telefonos"]
            direccion = alumno["direccion_residencia"]
            especialidad_ocupacional = alumno["especialidad_ocupacional"]
            docentes_evaluadores = alumno["lista_docentes_evaluadores"]
            fecha_evaluacion = alumno["fecha_evaluacion"]
            anio_escolar = alumno["anio_escolar"]
            
            
            # Título y subtítulo
            parrafo = documento.add_paragraph()
            parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = parrafo.add_run("INFORME EDUCATIVO INTEGRAL")
            run.bold = True
            self.cargar_fuente(run, "Arial", 12)
            
            self.agregar_salto_linea(documento, 1)
            
            parrafo = documento.add_paragraph()
            run = parrafo.add_run("Datos de Identificación:")
            run.bold = True
            self.cargar_fuente(run, "Arial", 12)
            
            self.agregar_salto_linea(documento, 1)
            
            # Nombre completo del Alumno
            parrafo = documento.add_paragraph()
            run_campo = parrafo.add_run("Nombres y Apellidos: ")
            run_campo.bold = True
            self.cargar_fuente(run_campo, "Arial", 11)
            
            run_valor = parrafo.add_run(nombre_completo_alumno)
            self.cargar_fuente(run_valor, "Arial", 11)
            
            
            # Cédula del Alumno
            parrafo = documento.add_paragraph()
            run_campo = parrafo.add_run("Cédula de Identidad: ")
            run_campo.bold = True
            self.cargar_fuente(run_campo, "Arial", 11)
            
            run_valor = parrafo.add_run(cedula_alumno)
            self.cargar_fuente(run_valor, "Arial", 11)
            
            
            # Lugar y Fecha de nacimiento
            parrafo = documento.add_paragraph()
            run_campo = parrafo.add_run("Lugar y Fecha de Nacimiento: ")
            run_campo.bold = True
            self.cargar_fuente(run_campo, "Arial", 11)
            
            run_valor = parrafo.add_run(lugar_y_fecha_nacimiento)
            self.cargar_fuente(run_valor, "Arial", 11)
            
            
            # Edad cronológica
            parrafo = documento.add_paragraph()
            run_campo = parrafo.add_run("Edad Cronológica: ")
            run_campo.bold = True
            self.cargar_fuente(run_campo, "Arial", 11)
            
            run_valor = parrafo.add_run(edad)
            self.cargar_fuente(run_valor, "Arial", 11)
            
            
            # Sexo
            parrafo = documento.add_paragraph()
            run_campo = parrafo.add_run("Sexo: ")
            run_campo.bold = True
            self.cargar_fuente(run_campo, "Arial", 11)
            
            run_valor = parrafo.add_run(sexo)
            self.cargar_fuente(run_valor, "Arial", 11)
            
            
            # Procedencia
            parrafo = documento.add_paragraph()
            run_campo = parrafo.add_run("Procedencia: ")
            run_campo.bold = True
            self.cargar_fuente(run_campo, "Arial", 11)
            
            run_valor = parrafo.add_run(procedencia)
            self.cargar_fuente(run_valor, "Arial", 11)
            
            
            # Diagnóstico
            parrafo = documento.add_paragraph()
            run_campo = parrafo.add_run("Diagnóstico: ")
            run_campo.bold = True
            self.cargar_fuente(run_campo, "Arial", 11)
            
            for diagnostico in diagnosticos_con_obervacion:
                condicion = diagnostico[0]
                medico_tratante = diagnostico[1]
                fecha_diagnostico = diagnostico[2]
                observacion_adicional = diagnostico[3]
                
                run_valor = parrafo.add_run(f"Informe realizado por {medico_tratante} en {fecha_diagnostico}. Presenta la condición de {condicion}. {observacion_adicional}. ")
                self.cargar_fuente(run_valor, "Arial", 11)
            
            
            # Medicación
            parrafo = documento.add_paragraph()
            run_campo = parrafo.add_run("Medicación: ")
            run_campo.bold = True
            self.cargar_fuente(run_campo, "Arial", 11)
            
            run_valor = parrafo.add_run(medicacion)
            self.cargar_fuente(run_valor, "Arial", 11)
            
            
            # Certificado de Discapacidad
            parrafo = documento.add_paragraph()
            run_campo = parrafo.add_run("Certificado de Discapacidad: ")
            run_campo.bold = True
            self.cargar_fuente(run_campo, "Arial", 11)
            
            run_valor = parrafo.add_run(certificado_discapacidad)
            self.cargar_fuente(run_valor, "Arial", 11)
            
            
            # Fecha de Ingreso al Taller
            parrafo = documento.add_paragraph()
            run_campo = parrafo.add_run("Fecha de Ingreso al Taller: ")
            run_campo.bold = True
            self.cargar_fuente(run_campo, "Arial", 11)
            
            run_valor = parrafo.add_run(fecha_ingreso_taller)
            self.cargar_fuente(run_valor, "Arial", 11)
            
            
            # Fecha de Ingreso a la Especialidad
            parrafo = documento.add_paragraph()
            run_campo = parrafo.add_run("Fecha de Ingreso a la Especialidad: ")
            run_campo.bold = True
            self.cargar_fuente(run_campo, "Arial", 11)
            
            run_valor = parrafo.add_run(fecha_ingreso_especialidad)
            self.cargar_fuente(run_valor, "Arial", 11)
            
            
            # Nombre y Apellido del Representante
            parrafo = documento.add_paragraph()
            run_campo = parrafo.add_run("Nombre y Apellido del Representante: ")
            run_campo.bold = True
            self.cargar_fuente(run_campo, "Arial", 11)
            
            run_valor = parrafo.add_run(nombre_completo_representante)
            self.cargar_fuente(run_valor, "Arial", 11)
            
            
            # Cédula del Representante
            parrafo = documento.add_paragraph()
            run_campo = parrafo.add_run("Cédula de Identidad del Representante: ")
            run_campo.bold = True
            self.cargar_fuente(run_campo, "Arial", 11)
            
            run_valor = parrafo.add_run(cedula_representante)
            self.cargar_fuente(run_valor, "Arial", 11)
            
            
            # Teléfono de Contacto
            parrafo = documento.add_paragraph()
            run_campo = parrafo.add_run("Teléfono de Contacto: ")
            run_campo.bold = True
            self.cargar_fuente(run_campo, "Arial", 11)
            
            run_valor = parrafo.add_run(telefono_representante)
            self.cargar_fuente(run_valor, "Arial", 11)
            
            
            # Dirección
            parrafo = documento.add_paragraph()
            run_campo = parrafo.add_run("Dirección: ")
            run_campo.bold = True
            self.cargar_fuente(run_campo, "Arial", 11)
            
            run_valor = parrafo.add_run(direccion)
            self.cargar_fuente(run_valor, "Arial", 11)
            
            
            # Especialidad Ocupacional
            parrafo = documento.add_paragraph()
            run_campo = parrafo.add_run("Especialidad Ocupacional: ")
            run_campo.bold = True
            self.cargar_fuente(run_campo, "Arial", 11)
            
            run_valor = parrafo.add_run(especialidad_ocupacional)
            self.cargar_fuente(run_valor, "Arial", 11)
            
            
            # Docentes Evaluadores
            parrafo = documento.add_paragraph()
            run_campo = parrafo.add_run("Docente Evaluador: ")
            run_campo.bold = True
            self.cargar_fuente(run_campo, "Arial", 11)
            
            run_valor = parrafo.add_run(docentes_evaluadores)
            self.cargar_fuente(run_valor, "Arial", 11)
            
            
            # Fecha de la Evaluación
            parrafo = documento.add_paragraph()
            run_campo = parrafo.add_run("Fecha de la Evaluación: ")
            run_campo.bold = True
            self.cargar_fuente(run_campo, "Arial", 11)
            
            run_valor = parrafo.add_run(fecha_evaluacion)
            self.cargar_fuente(run_valor, "Arial", 11)
            
            
            # Año Escolar
            parrafo = documento.add_paragraph()
            run_campo = parrafo.add_run("Año Escolar: ")
            run_campo.bold = True
            self.cargar_fuente(run_campo, "Arial", 11)
            
            run_valor = parrafo.add_run(anio_escolar)
            self.cargar_fuente(run_valor, "Arial", 11)
            
            
            self.agregar_salto_linea(documento, 1)
            self.agregar_salto_linea(documento, 1)
            self.agregar_salto_linea(documento, 1)
            
            
            # Área Personal Social
            parrafo = documento.add_paragraph().add_run("Área Personal Social:")
            parrafo.bold = True
            self.cargar_fuente(parrafo, "Arial", 12)
            
            self.agregar_salto_linea(documento, 1)
            
            # Área Académica
            parrafo = documento.add_paragraph().add_run("Área Académica:")
            parrafo.bold = True
            self.cargar_fuente(parrafo, "Arial", 12)
            
            self.agregar_salto_linea(documento, 1)
            
            # Área Laboral
            parrafo = documento.add_paragraph().add_run("Área Laboral:")
            parrafo.bold = True
            self.cargar_fuente(parrafo, "Arial", 12)
            
            self.agregar_salto_linea(documento, 1)
            
            # Conclusiones
            parrafo = documento.add_paragraph().add_run("Conclusiones:")
            parrafo.bold = True
            self.cargar_fuente(parrafo, "Arial", 12)
            
            self.agregar_salto_linea(documento, 1)
            
            # Recomendaciones
            parrafo = documento.add_paragraph().add_run("Recomendaciones:")
            parrafo.bold = True
            self.cargar_fuente(parrafo, "Arial", 12)
            
            self.agregar_salto_linea(documento, 1)
            
            # Cargar las firmas de los conformes
            self.cargar_espacio_firmas(documento, lista_dict_firmantes)
            
            self.agregar_salto_linea(documento, 1)
            
            documento.add_page_break()
        
    def cargar_datos(
        self, alumno_repositorio: AlumnoRepositorio,
        inscripcion_repositorio: InscripcionRepositorio,
        info_clinica_alumno_repositorio: InfoClinicaAlumnoRepositorio,
        detalles_cargo_repositorio: DetalleCargoRepositorio,
        especialidad_repositorio: EspecialidadRepositorio,
        especialidad_id: int,
        nombre_completo_coord_academico: str
    ) -> Optional[List[Any]]:
        try:
            datos = []
            
            especialidad_servicio = EspecialidadServicio(especialidad_repositorio)
            
            nombre_completo_directora = self.obtener_nombre_directora(detalles_cargo_repositorio)
            nombre_especialidad_ocupacional = especialidad_servicio.obtener_especialidad_por_id(especialidad_id)[1]
            
            lista_evaluadores_firmantes = self.transformar_lista_evaluadores(
                detalles_cargo_repositorio, especialidad_id
            )
            
            lista_dict_firmantes = self.transformar_lista_firmantes(
                lista_evaluadores_firmantes, nombre_completo_directora,
                nombre_completo_coord_academico
            )
            
            lista_evaluadores_especialidad = self.transformar_lista_evaluadores(
                detalles_cargo_repositorio, especialidad_id
            )
            
            lista_dict_alumnos = self.transformar_lista_data_alumnos(
                alumno_repositorio, inscripcion_repositorio,
                info_clinica_alumno_repositorio, especialidad_id,
                lista_evaluadores_especialidad
            )
            
            
            # Posición 0: Nombre de la especialidad ocupacional
            datos.append(nombre_especialidad_ocupacional)
            
            # Posición 1: Lista diccionario de la data de alumnos
            datos.append(lista_dict_alumnos)
            
            # Posición 2: Lista diccionario de los firmantes
            datos.append(lista_dict_firmantes)
            
            return datos
        except BaseDatosError as error:
            raise error
        except Exception as error:
            print(f"ERROR AL CARGAR LOS DATOS: {error}")
    
    def exportar(self, datos: List):
        try:
            self.RUTA_REPORTES_INFORMES_EDUCATIVOS_ALUMNOS.mkdir(exist_ok = True)
            
            nombre_especialidad_ocupacional = datos[0].upper()
            mes_actual = self.cargar_mes().upper()
            lista_dict_alumnos = datos[1]
            lista_dict_firmantes = datos[2]
            
            documento = self.crear_documento()
            fecha_actual = int(datetime.now().date().year)
            
            self.cargar_cintillo(documento)
            self.cargar_info_alumnos(documento, lista_dict_alumnos, lista_dict_firmantes)
            
            nombre_archivo = f"INFORME EDUCATIVO INTEGRAL DE {nombre_especialidad_ocupacional} {mes_actual}-{fecha_actual}"
            
            documento.save(f"{self.RUTA_REPORTES_INFORMES_EDUCATIVOS_ALUMNOS}/{nombre_archivo}.docx")
            
            documento_final = Document(f"{self.RUTA_REPORTES_INFORMES_EDUCATIVOS_ALUMNOS}/{nombre_archivo}.docx")
    
            if (len(documento_final.paragraphs) > 0):
                ultimo_parrafo = documento_final.paragraphs[-1]
                
                if (ultimo_parrafo.text.strip() == ""):
                    documento_final._element.body.remove(ultimo_parrafo._element)
    
            documento_final.save(f"{self.RUTA_REPORTES_INFORMES_EDUCATIVOS_ALUMNOS}/{nombre_archivo}.docx")
        except Exception as error:
            raise error


if __name__ == "__main__":
    reporte_informe_educativo_alumnos = ReporteInformeEducativoAlumnos()
    
    alumno_repositorio = AlumnoRepositorio()
    inscripcion_repositorio = InscripcionRepositorio()
    info_clinica_alumno_repositorio = InfoClinicaAlumnoRepositorio()
    detalles_cargo_repositorio = DetalleCargoRepositorio()
    especialidad_repositorio = EspecialidadRepositorio()
    
    ESPECIALIDAD_ID = 1
    nombre_completo_coord_academico = "FULANO DE TAL"
    
    datos = reporte_informe_educativo_alumnos.cargar_datos(
        alumno_repositorio,
        inscripcion_repositorio,
        info_clinica_alumno_repositorio,
        detalles_cargo_repositorio,
        especialidad_repositorio,
        ESPECIALIDAD_ID,
        nombre_completo_coord_academico
    )
    
    reporte_informe_educativo_alumnos.exportar(datos)
