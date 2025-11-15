from typing import List, Dict, Optional, Any
from datetime import datetime
from configuraciones.configuracion import app_configuracion
from reportes.reporte_base import ReporteBase
from excepciones.base_datos_error import BaseDatosError

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from repositorios.alumnos.alumno_repositorio import AlumnoRepositorio
from repositorios.alumnos.medidas_alumno_repositorio import MedidasAlumnoRepositorio
from repositorios.alumnos.inscripcion_repositorio import InscripcionRepositorio
from repositorios.alumnos.info_clinica_alumno_repositorio import InfoClinicaAlumnoRepositorio

from servicios.alumnos.alumno_servicio import AlumnoServicio
from servicios.alumnos.medidas_alumno_servicio import MedidasAlumnoServicio
from servicios.alumnos.inscripcion_servicio import InscripcionServicio
from servicios.alumnos.info_clinica_alumno_servicio import InfoClinicaAlumnoServicio


class ReporteGeneralAlumnos(ReporteBase):
    def __init__(self):
        self.RUTA_REPORTES_GENERALES_ALUMNOS = app_configuracion.DIRECTORIO_REPORTES_ALUMNOS
    
    def transformar_data_alumnos(
        self, inscripcion_repositorio: InscripcionRepositorio,
        alumno_repositorio: AlumnoRepositorio,
        medidas_alumno_repositorio: MedidasAlumnoRepositorio,
        info_clinica_alumno_repositorio: InfoClinicaAlumnoRepositorio
    ) -> List[Dict]:
        lista_data_alumnos = []
        
        alumno_servicio = AlumnoServicio(alumno_repositorio)
        medidas_alumno_servicio = MedidasAlumnoServicio(medidas_alumno_repositorio)
        inscripcion_servicio = InscripcionServicio(inscripcion_repositorio)
        info_clinica_alumno_servicio = InfoClinicaAlumnoServicio(info_clinica_alumno_repositorio)
        
        for num, alumno in enumerate(alumno_servicio.obtener_todos_alumnos(), start = 1):
            num_registro = num
            alumno_id = alumno[0]
            
            primer_nombre = alumno[2]
            segundo_nombre = alumno[3] if (alumno[3]) else ""
            tercer_nombre = alumno[4] if (alumno[4]) else ""
            apellido_paterno = alumno[5]
            apellido_materno = alumno[6] if (alumno[6]) else ""
            nombre_completo = f"{primer_nombre} {segundo_nombre} {tercer_nombre} {apellido_paterno} {apellido_materno}"
            
            cedula = alumno[1]
            
            especialidad_ocupacional = inscripcion_servicio.obtener_inscripcion_por_id(alumno_id)[3]
            
            lugar_nacimiento = alumno[9]
            fecha_nacimiento = alumno[7]
            lugar_y_fecha_nacimiento = f"{lugar_nacimiento} {fecha_nacimiento}"
            
            edad = alumno[8]
            sexo = alumno[10]
            
            lista_diagnosticos = []
            lista_medicacion = []
            lista_certificado_discapacidad = []
            
            
            for diagnostico in info_clinica_alumno_servicio.obtener_info_clinica_por_alumno_id(alumno_id):
                nombre_diagnostico = diagnostico[2] if (diagnostico[2] is not None) else ""
                medicacion = diagnostico[7] if (diagnostico[7] is not None) else ""
                certificado_discapacidad = diagnostico[5] if (diagnostico[5] is not None) else ""
                
                lista_diagnosticos.append(nombre_diagnostico)
                lista_medicacion.append(medicacion)
                lista_certificado_discapacidad.append(certificado_discapacidad)
            
            talla = medidas_alumno_servicio.obtener_medidas_alumno_por_id(alumno_id)[2]
            peso = medidas_alumno_servicio.obtener_medidas_alumno_por_id(alumno_id)[3]
            talla_camisa = medidas_alumno_servicio.obtener_medidas_alumno_por_id(alumno_id)[4]
            talla_pantalon = medidas_alumno_servicio.obtener_medidas_alumno_por_id(alumno_id)[5]
            talla_zapatos = medidas_alumno_servicio.obtener_medidas_alumno_por_id(alumno_id)[6]
            
            procedencia = alumno_servicio.obtener_info_academica_alumno(alumno_id)[2]
            fecha_ingreso_institucion = alumno[13]
            escolaridad = alumno_servicio.obtener_info_academica_alumno(alumno_id)[1]
            tiempo_en_tela = inscripcion_servicio.obtener_inscripcion_por_id(alumno_id)[7]
            cma = alumno[11]
            imt = alumno[12]
            
            nombre_representante = alumno_servicio.obtener_datos_representante(alumno_id)[3]
            apellido_representante = alumno_servicio.obtener_datos_representante(alumno_id)[4]
            nombre_completo_representante = f"{nombre_representante} {apellido_representante}"
            cedula_representante = alumno_servicio.obtener_datos_representante(alumno_id)[2]
            direccion_residencia = alumno_servicio.obtener_datos_representante(alumno_id)[5]
            num_telefono_principal = alumno_servicio.obtener_datos_representante(alumno_id)[6]
            num_telefono_secundario = alumno_servicio.obtener_datos_representante(alumno_id)[7] if (alumno_servicio.obtener_datos_representante(alumno_id)[7] is not None) else ""
            lista_telefonos = [num_telefono_principal, num_telefono_secundario]
            
            if (len(lista_telefonos) == 2):
                telefonos = "\n".join(lista_telefonos)
            else:
                telefonos = num_telefono_principal
            
            carga_familiar = alumno_servicio.obtener_datos_representante(alumno_id)[8]
            
            elemento_lista = {
                "num": num_registro,
                "nombre_completo_alumno": nombre_completo,
                "cedula": cedula,
                "especialidad_ocupacional": especialidad_ocupacional,
                "lugar_y_fecha_nacimiento": lugar_y_fecha_nacimiento,
                "edad": edad,
                "sexo": sexo,
                "lista_diagnosticos": "\n".join(lista_diagnosticos),
                "lista_medicacion": "\n".join(lista_medicacion),
                "talla": talla,
                "peso": peso,
                "talla_camisa": talla_camisa,
                "talla_pantalon": talla_pantalon,
                "talla_zapatos": talla_zapatos,
                "procedencia": procedencia,
                "fecha_ingreso_institucion": fecha_ingreso_institucion,
                "escolaridad": escolaridad,
                "tiempo_en_tela": tiempo_en_tela,
                "cma": cma,
                "imt": imt,
                "lista_certificado_discapacidad": "\n".join(lista_certificado_discapacidad),
                "nombre_completo_representante": nombre_completo_representante,
                "cedula_representante": cedula_representante,
                "direccion_residencia": direccion_residencia,
                "telefonos": telefonos,
                "carga_familiar": carga_familiar
            }
            
            lista_data_alumnos.append(elemento_lista)
        
        return lista_data_alumnos
    
    def cargar_datos(
            self, inscripcion_repositorio: InscripcionRepositorio,
            alumno_repositorio: AlumnoRepositorio,
            medidas_alumno_repositorio: MedidasAlumnoRepositorio,
            info_clinica_alumno_repositorio: InfoClinicaAlumnoRepositorio
        ) -> Optional[List[Any]]:
        try:
            datos = []
            
            # Posición 0: Lista de toda la data de alumnos
            lista_dict_data_alumnos = self.transformar_data_alumnos(
                inscripcion_repositorio, alumno_repositorio,
                medidas_alumno_repositorio, info_clinica_alumno_repositorio
            )
            
            datos.append(lista_dict_data_alumnos)
            
            return datos
        except BaseDatosError as error:
            raise error
        except Exception as error:
            print(f"ERROR AL CARGAR LOS DATOS: {error}")
    
    def crear_documento(self):
        documento = Document()
        
        return documento
    
    def cargar_configuraciones_excel(self, documento):
        # Para cambiar la orientación del documento a Horizontal
        seccion = documento.sections[0]


        # Configurar el tamaño de la hoja a Legal
        seccion.page_width = Cm(21.59)
        seccion.page_height = Cm(35.56)

        # Intercambiamos el ancho y el alto para cambiar a Horizontal
        nuevo_ancho, nuevo_alto = seccion.page_height, seccion.page_width

        seccion.page_width = nuevo_ancho
        seccion.page_height = nuevo_alto

        # Asignar la orientación a LANDSCAPE
        seccion.orientation = WD_ORIENT.LANDSCAPE

        # Cambiar los márgenes para que cada lado tenga 1.27cm
        seccion.left_margin = Cm(1.27)
        seccion.right_margin = Cm(1.27)
        seccion.top_margin = Cm(1.27)
        seccion.bottom_margin = Cm(1.27)
    
    def cargar_cintillo(self, documento):
        encabezado = documento.sections[0].header

        encabezado_parrafo = encabezado.paragraphs[0] if encabezado.paragraphs else encabezado.add_paragraph()
        encabezado_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = encabezado_parrafo.add_run()
        run.add_picture("reportes/imagenes/CINTILLO_GRANDE.png")
    
    def cargar_titulo(self, documento):
        # Cargar el título
        parrafo = documento.add_paragraph()
        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        fecha_actual = int(datetime.now().date().year)
        
        run = parrafo.add_run(f"CARACTERIZACIÓN DE LOS ALUMNOS DEL TALLER DE EDUCACIÓN LABORAL ANZOÁTEGUI {fecha_actual}-{fecha_actual + 1}")
        run.font.name = "Calibri"
        run.font.size = Pt(12)
        run.bold = True
        run.underline = True
        run.add_break()
    
    def cargar_data_alumnos(self, documento, lista_dict_data_alumnos: List[Dict]):
        for i in range(0, len(lista_dict_data_alumnos), 8):
            # Obtenemos un trozo de 15 personas para esta iteración
            bloque_personas = lista_dict_data_alumnos[i:i + 8]
            
            # Tabla 1 (Información principal)
            tabla_1 = documento.add_table(rows = 1, cols = 16)
            tabla_1.style = "Table Grid"
            
            encabezado_celdas_1 = tabla_1.rows[0].cells
            encabezado_celdas_1[0].text = "N"
            encabezado_celdas_1[1].text = "APELLIDOS Y NOMBRES"
            encabezado_celdas_1[2].text = "C.I"
            encabezado_celdas_1[3].text = "ESP. OCUP."
            encabezado_celdas_1[4].text = "LUGAR Y F/N"
            encabezado_celdas_1[5].text = "EDAD"
            encabezado_celdas_1[6].text = "SEXO"
            encabezado_celdas_1[7].text = "DIAGNÓSTICO"
            encabezado_celdas_1[8].text = "MEDICACIÓN"
            encabezado_celdas_1[9].text = "TALLA"
            encabezado_celdas_1[10].text = "PESO"
            encabezado_celdas_1[11].text = "C"
            encabezado_celdas_1[12].text = "P"
            encabezado_celdas_1[13].text = "Z"
            encabezado_celdas_1[14].text = "PROCEDENCIA"
            encabezado_celdas_1[15].text = "F. INGRESO"
            
            # Establecer anchos de la columna fijos en Cm de la tabla 1
            tabla_1.columns[0].width = Cm(0.3)
            tabla_1.columns[1].width = Cm(4.5)
            tabla_1.columns[2].width = Cm(2.5)
            tabla_1.columns[3].width = Cm(2.0)
            tabla_1.columns[4].width = Cm(3.5)
            tabla_1.columns[5].width = Cm(1.0)
            tabla_1.columns[6].width = Cm(1.0)
            tabla_1.columns[7].width = Cm(4.0)
            tabla_1.columns[8].width = Cm(4.0)
            tabla_1.columns[9].width = Cm(1.0)
            tabla_1.columns[10].width = Cm(1.0)
            tabla_1.columns[11].width = Cm(0.2)
            tabla_1.columns[12].width = Cm(0.2)
            tabla_1.columns[13].width = Cm(0.2)
            tabla_1.columns[14].width = Cm(2.5)
            tabla_1.columns[15].width = Cm(2.5)
            
            # Configurar el estilo de los encabezados de la tabla 1
            for celda in encabezado_celdas_1:
                celda.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in celda.paragraphs[0].runs:
                    run.font.name = "Calibri"
                    run.font.bold = True
                    run.font.size = Pt(11)
            
            
            # Llenar la primera tabla con los datos del bloque actual
            for persona in bloque_personas:
                fila_celdas = tabla_1.add_row().cells
                fila_celdas[0].text = str(persona["num"])
                fila_celdas[1].text = persona["nombre_completo_alumno"]
                fila_celdas[2].text = persona["cedula"]
                fila_celdas[3].text = persona["especialidad_ocupacional"]
                fila_celdas[4].text = persona["lugar_y_fecha_nacimiento"]
                fila_celdas[5].text = str(persona["edad"])
                fila_celdas[6].text = persona["sexo"]
                fila_celdas[7].text = persona["lista_diagnosticos"]
                fila_celdas[8].text = persona["lista_medicacion"]
                fila_celdas[9].text = str(persona["talla"])
                fila_celdas[10].text = str(persona["peso"])
                fila_celdas[11].text = persona["talla_camisa"]
                fila_celdas[12].text = str(persona["talla_pantalon"])
                fila_celdas[13].text = str(persona["talla_zapatos"])
                fila_celdas[14].text = persona["procedencia"]
                fila_celdas[15].text = persona["fecha_ingreso_institucion"]
                
                # Configurar el estilo al contenido de cada fila de la tabla 1
                for celda in fila_celdas:
                    celda.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in celda.paragraphs[0].runs:
                        run.font.name = "Calibri"
                        run.font.bold = False
                        run.font.size = Pt(11)
            
            
            # Tabla 2 (Información secundaria)
            # Después de la primera tabla, insertamos un salto de página
            documento.add_page_break()
            
            tabla_2 = documento.add_table(rows = 1, cols = 11)
            tabla_2.style = "Table Grid"
            
            encabezado_celdas_2 = tabla_2.rows[0].cells
            encabezado_celdas_2[0].text = "N"
            encabezado_celdas_2[1].text = "ESCOLARIDAD"
            encabezado_celdas_2[2].text = "TIEMPO EN EL TELA"
            encabezado_celdas_2[3].text = "C.M.A"
            encabezado_celdas_2[4].text = "I.M.T"
            encabezado_celdas_2[5].text = "CERTIFICADO DE DISCAPACIDAD"
            encabezado_celdas_2[6].text = "REPRESENTANTE"
            encabezado_celdas_2[7].text = "C.I"
            encabezado_celdas_2[8].text = "DIRECCIÓN"
            encabezado_celdas_2[9].text = "TELÉFONOS"
            encabezado_celdas_2[10].text = "CARGA FAMILIAR"
            
            # Establecer anchos de la columna fijos en Cm de la tabla 2
            tabla_2.columns[0].width = Cm(0.3)
            tabla_2.columns[1].width = Cm(3.0)
            tabla_2.columns[2].width = Cm(3.0)
            tabla_2.columns[3].width = Cm(1.2)
            tabla_2.columns[4].width = Cm(1.2)
            tabla_2.columns[5].width = Cm(3.5)
            tabla_2.columns[6].width = Cm(5.0)
            tabla_2.columns[7].width = Cm(2.5)
            tabla_2.columns[8].width = Cm(6.0)
            tabla_2.columns[9].width = Cm(2.5)
            tabla_2.columns[10].width = Cm(1.5)
            
            # Configurar el estilo de los encabezados de la tabla 1
            for celda in encabezado_celdas_2:
                celda.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in celda.paragraphs[0].runs:
                    run.font.name = "Calibri"
                    run.font.bold = True
                    run.font.size = Pt(11)
            
            # Llenar la segunda tabla con los datos del mismo bloque
            for persona in bloque_personas:
                tiempo_en_tela = persona["tiempo_en_tela"]
                
                fila_celdas = tabla_2.add_row().cells
                fila_celdas[0].text = str(persona["num"])
                fila_celdas[1].text = persona["escolaridad"]
                fila_celdas[2].text = f"{tiempo_en_tela} años" if (tiempo_en_tela > 1) else "Nuevo ingreso"
                fila_celdas[3].text = "No" if (persona["cma"] == 0) else "Si"
                fila_celdas[4].text = "No" if (persona["imt"] == 0) else "Si"
                fila_celdas[5].text = persona["lista_certificado_discapacidad"]
                fila_celdas[6].text = persona["nombre_completo_representante"]
                fila_celdas[7].text = persona["cedula_representante"]
                fila_celdas[8].text = persona["direccion_residencia"]
                fila_celdas[9].text = persona["telefonos"]
                fila_celdas[10].text = str(persona["carga_familiar"])
                
                # Configurar el estilo al contenido de cada fila de la tabla 2
                for celda in fila_celdas:
                    celda.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in celda.paragraphs[0].runs:
                        run.font.name = "Calibri"
                        run.font.bold = False
                        run.font.size = Pt(11)
            
            
            # Acá se inserta otro salto de página para que el siguiente par de tablas comience en una nueva hoja.
            # Se puede hacer esto en cada iteración excepto en la última
            if i + 8 < len(lista_dict_data_alumnos):
                documento.add_page_break()
    
    def exportar(self, datos: List):
        try:
            self.RUTA_REPORTES_GENERALES_ALUMNOS.mkdir(exist_ok = True)
            
            lista_dict_data_alumnos = datos[0]
            
            # Crear el documento
            documento = self.crear_documento()
            
            # Cargar las configuraciones del documento
            self.cargar_configuraciones_excel(documento)
            
            # Cargar el cintillo
            self.cargar_cintillo(documento)
            
            # Cargar el título
            self.cargar_titulo(documento)
            
            # Nombre del documento
            fecha_actual = int(datetime.now().date().year)
            nombre_archivo = f"CARACTERIZACIÓN GENERAL TELA {fecha_actual}-{fecha_actual + 1}"
            
            # Cargar toda la data de todos los alumnos
            self.cargar_data_alumnos(documento, lista_dict_data_alumnos)
            
            # Guardar el documento
            documento.save(f"{self.RUTA_REPORTES_GENERALES_ALUMNOS}/{nombre_archivo}.docx")
        except Exception as error:
            raise error



if __name__ == "__main__":
    reporte_general_alumnos = ReporteGeneralAlumnos()
    
    inscripcion_repositorio = InscripcionRepositorio()
    alumno_repositorio = AlumnoRepositorio()
    medidas_alumno_repositorio = MedidasAlumnoRepositorio()
    info_clinica_alumno_repositorio = InfoClinicaAlumnoRepositorio()
    
    datos = reporte_general_alumnos.cargar_datos(
        inscripcion_repositorio, alumno_repositorio,
        medidas_alumno_repositorio, info_clinica_alumno_repositorio
    )
    
    reporte_general_alumnos.exportar(datos)
